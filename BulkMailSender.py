# -*- coding: utf-8 -*-

from importlib.resources import path
from sys import path_hooks
import pypandoc

from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtWidgets import QFileDialog ,QLabel

import string
from docx import Document
import pandas as pd
from docx2pdf import convert
import smtplib
from email.mime.multipart import MIMEMultipart
from email.message import EmailMessage

from email.mime.text import MIMEText
from email.mime.base import MIMEBase

from PyQt6.QtWidgets import QMessageBox
from email import encoders
import os
import time
from email.header import Header
from email.utils import formataddr
from pypandoc.pandoc_download import download_pandoc
from PyPDF2 import PdfMerger

class Ui_MainWindow(object):
    def setupUi(self, MainWindow,btn_val,gb_val,title_name,mail_subject,mail_body,all_text_color,send_button_text_color,wait_time,main_window_title,main_icon):
        self.mail_body = mail_body
        self.mail_subject= mail_subject
        self.all_text_color_val=all_text_color
        self.all_text_color="font: 75 13pt \"MS Shell Dlg 2\";color:rgb%s;"%(self.all_text_color_val)
        self.all_text_color_browse="color:rgb%s;"%(self.all_text_color_val)
        self.send_button_text_color = send_button_text_color
        self.wait_time = wait_time
        self.background_color_value = gb_val
        self.background_color = "background-color: rgb%s;"%(self.background_color_value)


        self.button_color_value = btn_val

        self.button_color = "background-color: rgb%s;color: rgb%s"%(self.button_color_value,self.send_button_text_color)
        
        
        # MainWindow.setObjectName("Email sender")
        
        MainWindow.resize(562, 600)
        MainWindow.setWindowTitle(main_window_title)
        MainWindow.setWindowIcon(QtGui.QIcon(main_icon))

    
        
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        # self.windowTitle("New Title")
        # self.setWindowIcon(QtGui.QIcon('123.png'))
        

        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.frame = QtWidgets.QFrame(self.centralwidget)
        
        font = QtGui.QFont()
        font.setPointSize(14)
        self.frame.setFont(font)
        self.frame.setStyleSheet(self.background_color)
        self.frame.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Shadow.Raised)
        self.frame.setObjectName("frame")
        self.toolButton = QtWidgets.QToolButton(self.frame)
        self.toolButton.setGeometry(QtCore.QRect(400, 110, 61, 31))
        self.toolButton.setObjectName("toolButton")
        self.toolButton.clicked.connect(self.exl_path)
        self.toolButton.setStyleSheet(self.all_text_color_browse)

        self.pushButton = QtWidgets.QPushButton(self.frame)
        self.pushButton.setGeometry(QtCore.QRect(240, 470, 121, 51))
        self.pushButton.setStyleSheet(self.button_color)
        self.pushButton.setObjectName("pushButton")
        self.label = QtWidgets.QLabel(self.frame)
        self.label.setGeometry(QtCore.QRect(80, 110, 310, 31))
        font = QtGui.QFont()
        font.setFamily("MS Shell Dlg 2")
        font.setPointSize(13)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(9)
        self.label.setFont(font)

        self.label.setStyleSheet(self.all_text_color)
        self.label.setObjectName("label")
        
        self.label_2 = QtWidgets.QLabel(self.frame)
        self.label_2.setGeometry(QtCore.QRect(80, 160, 310, 31))
        self.label_2.setStyleSheet(self.all_text_color)
        self.label_2.setObjectName("label_2")
        self.label_4 = QtWidgets.QLabel(self.frame)
        self.label_4.setStyleSheet(self.all_text_color)
        self.label_4.setText(title_name)
        self.label_4.setGeometry(QtCore.QRect(220, 20, 241, 61))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.comboBox = QtWidgets.QComboBox(self.frame)
        self.comboBox.setGeometry(QtCore.QRect(330, 340, 171, 21))
        self.comboBox.setCurrentText("")
        self.comboBox.setObjectName("comboBox")
        self.comboBox.setStyleSheet(self.all_text_color_browse)

        self.toolButton_2 = QtWidgets.QToolButton(self.frame)
        self.toolButton_2.setGeometry(QtCore.QRect(400, 160, 61, 31))
        self.toolButton_2.setObjectName("toolButton_2")
        self.toolButton_2.clicked.connect(self.output_path)
        self.toolButton_2.setStyleSheet(self.all_text_color_browse)

        self.label_5 = QtWidgets.QLabel(self.frame)
        self.label_5.setGeometry(QtCore.QRect(80, 210, 310, 31))
        # self.label_5.setGeometry(QtCore.QRect(80, 270, 231, 31))
        self.label_5.setStyleSheet(self.all_text_color)
        self.label_5.setObjectName("label_5")
        # self.toolButton_3 = QtWidgets.QToolButton(self.frame)
        # self.toolButton_3.setGeometry(QtCore.QRect(330, 210, 61, 31))
        # self.toolButton_3.setObjectName("toolButton_3")
        # self.toolButton_3.clicked.connect(self.attachment_template_path)
        # self.toolButton_3.setStyleSheet(self.all_text_color_browse)

        self.toolButton_4 = QtWidgets.QToolButton(self.frame)
        self.toolButton_4.setGeometry(QtCore.QRect(400, 210, 61, 31))
        # self.toolButton_4.setGeometry(QtCore.QRect(330, 270, 61, 31))
        self.toolButton_4.setObjectName("toolButton_4")
        self.toolButton_4.clicked.connect(self.body_template_path)
        self.toolButton_4.setStyleSheet(self.all_text_color_browse)



        # self.toolButton_5 = QtWidgets.QToolButton(self.frame)
        # # self.toolButton_5.setGeometry(QtCore.QRect(400, 270, 61, 31))
        # # self.toolButton_5.setGeometry(QtCore.QRect(400, 390, 61, 31))
        # self.toolButton_5.setObjectName("toolButton_5")
        # self.toolButton_5.clicked.connect(self.subject_template_path)
        # self.toolButton_5.setStyleSheet(self.all_text_color_browse)
        

        self.pushButton_2 = QtWidgets.QPushButton(self.frame)
        self.pushButton_2.setGeometry(QtCore.QRect(5, 540, 30, 30))
        self.pushButton_2.setText("")
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap("settings.png"), QtGui.QIcon.Mode.Normal, QtGui.QIcon.State.Off)
        self.pushButton_2.setIcon(icon)
        self.pushButton_2.setIconSize(QtCore.QSize(30, 30))
        self.pushButton_2.setCheckable(False)
        self.pushButton_2.setChecked(False)
        self.pushButton_2.setAutoDefault(False)
        self.pushButton_2.setDefault(False)
        self.pushButton_2.setObjectName("pushButton_2")
        self.pushButton_2.clicked.connect(self.ChangeSettings)


        self.label_7 = QtWidgets.QLabel(self.frame)
        self.label_7.setGeometry(QtCore.QRect(80, 270, 231, 31))
        # self.label_7.setGeometry(QtCore.QRect(80, 390, 310, 31))
        self.label_7.setStyleSheet(self.all_text_color)
        self.label_7.setObjectName("label_7")

        # self.lineEdit = QtWidgets.QLineEdit(self.frame)
        # self.lineEdit.setGeometry(QtCore.QRect(330, 390, 171, 31))
        # self.lineEdit.setStyleSheet(self.all_text_color_browse)
        # self.lineEdit.setObjectName("lineEdit")
        # self.lineEdit.setText(self.mail_subject)
        
        self.label_8 = QtWidgets.QLabel(self.frame)
        self.label_8.setGeometry(QtCore.QRect(120, 20, 91, 61))
        self.label_8.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.SizeBDiagCursor))
        self.label_8.setText("")
        self.label_8.setPixmap(QtGui.QPixmap(main_icon))
        self.label_8.setScaledContents(True)
        self.label_8.setObjectName("label_8")
        self.label_3 = QtWidgets.QLabel(self.frame)
        self.label_3.setGeometry(QtCore.QRect(80, 330, 141, 31))
        self.label_3.setStyleSheet(self.all_text_color)
        self.label_3.setObjectName("label_3")
        # self.label_6 = QtWidgets.QLabel(self.frame)
        # self.label_6.setGeometry(QtCore.QRect(80, 210, 161, 31))
        # self.label_6.setStyleSheet(self.all_text_color)
        # self.label_6.setObjectName("label_6")
        self.verticalLayout.addWidget(self.frame)
        MainWindow.setCentralWidget(self.centralwidget)

        #Subject Line
        self.lineEdit = QtWidgets.QLineEdit(self.frame)
        self.lineEdit.setGeometry(QtCore.QRect(330, 270, 171, 31))
        self.lineEdit.setStyleSheet(self.all_text_color_browse)
        self.lineEdit.setObjectName("lineEdit")
        self.lineEdit.setText("Mail From SC Legal")
        

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.exl_file_path = None
        self.output_folder_path = None
        self.body_tem_file_path= None
        self.subject_tem_file_path=None
        self.attachment_tem_file_path=None
       

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        
        
        self.toolButton.setText(_translate("MainWindow", "Browse"))
        # self.toolButton_3.setText(_translate("MainWindow", "Browse"))
        self.toolButton_4.setText(_translate("MainWindow", "Browse"))
        self.toolButton_2.setText(_translate("MainWindow", "Browse"))
        # self.toolButton_5.setText(_translate("MainWindow", "Browse"))

        self.pushButton.setText(_translate("MainWindow", "Send"))
        self.pushButton.clicked.connect(self.Send_button)
        
        self.label.setText(_translate("MainWindow", "Select Excel Sheet"))
        self.label_2.setText(_translate("MainWindow", "Select Output Folder"))
        # self.label_4.setText(_translate("MainWindow", "Email Notification Sender"))
        self.label_5.setText(_translate("MainWindow", "Select Attachment Template"))

        self.label_7.setText(_translate("MainWindow", "Mail Subject"))
        self.label_3.setText(_translate("MainWindow", "Select File Name"))
        # self.label_6.setText(_translate("MainWindow", "Select Body Template :"))

    def ChangeSettings(self):
        print("Changing Setting")
        # osCommandString = "notepad.exe Config_file.txt"
        # os.system(osCommandString)

        # os.startfile('Config_file.txt')
        import subprocess
        import platform as pf
        if sys.platform == "win32":
            subprocess.call(['notepad.exe', 'Config_file.txt'])
        elif sys.platform == "darwin":
            subprocess.call(['open', '-a', 'TextEdit', 'Config_file.txt'])
        else:
            self.Pop_up_message("System Not Suported")
            
    

    def body_template_path(self):
        print("selecting mail attachment file")
        files, _ = QFileDialog.getOpenFileName(None, "Open File", "", "Microsoft Word Documents (*.docx)")
        self.attachment_tem_file_path = str(files)
        print(self.attachment_tem_file_path)
        output = pypandoc.convert_file(self.attachment_tem_file_path, 'html', outputfile="AttachmentTempalate.html", encoding="utf-8")
        assert output == ""
        temp=""
        temp="AttachmentTempalate.html"
        # print(temp)
        files = temp
        self.body_tem_file_path = str(files)
        print(self.body_tem_file_path)


    # def attachment_template_path(self):
    #     print("selecting mail body Template ")
    #     files, _ = QFileDialog.getOpenFileName(None, "Open File", "", "HTML File (*.html)")
    #     self.body_tem_file_path = str(files)
    #     print(self.body_tem_file_path)

    def attachment_template_path(self):
        print("selecting mail body Template ")

    def subject_template_path(self):
        print("selecting mail body Template ")
        files, _ = QFileDialog.getOpenFileName(None, "Open File", "", "HTML File (*.html)")
        self.subject_tem_file_path = str(files)
        print(self.subject_tem_file_path)



    def exl_path(self):
        print("Exl path Button")
        files, _ = QFileDialog.getOpenFileName(None, "Open File", "", "Excel File (*.xlsx)")
        self.exl_file_path = str(files)
        value = self.comboBox.currentText()
        print(value)

        print(self.exl_file_path)
        
        if self.exl_file_path == "":
            print("No excel file selected.")
            return

        self.comboBox.clear()
        data = pd.read_excel(self.exl_file_path)
        
        all_columns = data.columns
        self.allColumns = all_columns
        self.comboBox.setCurrentText("")
        # self.lineEdit.setCurrentText
        for val in all_columns:
            self.comboBox.addItem(str(val))
        
        # self.lineEdit.setText(QFileDialog.getOpenFileName(None, "Open File", "Desktop", "Excel Workshee (*.xlsx)"))

    def output_path(self):
        print("output file button is clicked")
        file =QFileDialog.getExistingDirectory(None, "Select Folder")
        self.output_folder_path=str(file)
        
        print(self.output_folder_path)
        
    
    def mail_Body_message_formatted(self,body_template_html_path):
        try:
            all_text = ""
            with open(body_template_html_path,'r') as pdf:
                data = pdf.readlines()
                for line in data:
                    all_text = all_text+line
                # print(all_text)

                return all_text
        except Exception as e:
            print(e)
            print("Error in Email Body HTML Template.")
            data = "Please check the message below"
        return data

    def sendMail_new(self,pdf_file_to_send,send_to,body_template_html_path,mail_msg,subject_msg):

        # q=self.ch_var2(aaa,col_val,row_val)
        

        sender = lines[3].strip()
        sender = sender.split("=")[1]

        password = lines[4].strip()
        password = password.split("=")[1]

        receiver = send_to
        subj_tex=subject_msg

        msg = EmailMessage()
        msg['Subject'] = subj_tex
        msg['From'] = formataddr((str(Header('SC Legal', 'utf-8')), sender))
        msg['To'] = receiver

        smtp_host =lines[9].strip()
        smtp_host=smtp_host.split("=")[1]
       

        smtp_port =lines[10].strip()
        smtp_port=smtp_port.split("=")[1]
        smtp_port=int(smtp_port)

        pdfname = self.pdf_name_in_mail+".pdf"
        
        msg.set_content(mail_msg, subtype='html')

        with open(pdf_file_to_send, 'rb') as pdf:
            msg.add_attachment(pdf.read(), maintype='application', subtype='octet-stream', filename=pdfname)



        with smtplib.SMTP(smtp_host, smtp_port) as smtp:
            smtp.starttls()
            smtp.login(sender, password) 
            smtp.send_message(msg)

        print("Mail sent successfully.")

    def sendMail(self,pdf_file_to_send,send_to,user_name,col_val,row_val,path2):

        aaa=path2
        q=self.ch_var2(aaa,col_val,row_val)



        # body="akakakakak"
        body = q
        # path2.save("akaks.docx")
        # path2.close()
        # put your email here
        sender = lines[3].strip()
        sender = sender.split("=")[1]

        password = lines[4].strip()
        password = password.split("=")[1]

        receiver = send_to
        subj_tex=self.lineEdit.text()
        # print(tt)
        #Setup the MIME
        message = MIMEMultipart()
        message['From'] = sender
        message['To'] = receiver
        # message['Subject'] = 'Promotion Notification'
        message['Subject'] = subj_tex
        
        message.attach(MIMEText(body, 'plain'))
        
        pdfname = self.pdf_name_in_mail+".pdf"
        
        # open the file in bynary
        binary_pdf = open(pdf_file_to_send, 'rb')
        
        payload = MIMEBase('application', 'octate-stream', Name=pdfname)
        # payload = MIMEBase('application', 'pdf', Name=pdfname)
        payload.set_payload((binary_pdf).read())
        
        # enconding the binary into base64
        encoders.encode_base64(payload)
        
        # add header with pdf name
        payload.add_header('Content-Decomposition', 'attachment', filename=pdfname)
        message.attach(payload)


        smtp_host =lines[9].strip()
        smtp_host=smtp_host.split("=")[1]
       

        smtp_port =lines[10].strip()
        smtp_port=smtp_port.split("=")[1]
        smtp_port=int(smtp_port)
        


        #use gmail with port
        session = smtplib.SMTP(smtp_host, smtp_port)
        #enable security
        session.starttls()
        #login with mail_id and password
        session.login(sender, password)
        
        text = message.as_string()
        session.sendmail(sender, receiver, text)
        session.quit()
        print('Mail Sent')

    def ch_var(self, path, col_val, row_val):
    #start
        sample_page = path

        dic = {col_val[i]: row_val[i] for i in range(len(col_val))}
        print(dic)

        mail_all_body_text = self.mail_Body_message_formatted(self.body_tem_file_path)
        subject_text = self.mail_Body_message_formatted(self.subject_tem_file_path)
        for paragraph in sample_page.paragraphs:
            for i in range(len(paragraph.runs)):
                text = paragraph.runs[i].text
                for key in dic:
                    if str(key).upper() in str(text).upper():
                        text = text.replace(key, dic[key])
                        paragraph.runs[i].text = text
                    mail_all_body_text = mail_all_body_text.replace(key, dic[key])
                    subject_text = subject_text.replace(key, dic[key])
                    # print(key,", ", dic[key])

        return sample_page, mail_all_body_text, subject_text
    #end       


    def ch_var2(self,path,col_val,row_val):
        #start
        sample_page = path

        all_text = ""

        # for p in sample_page.paragraphs:
        #     print(p.text)
        #     all_text = all_text + p.text

        for paragraph in sample_page.paragraphs:
            for run in paragraph.runs:
            # print(run.text)
                if run.bold:
                    print(run.text)
                    start = ""
                    end = ""
                    all_text = all_text + start + run.text + end #
                else:
                    all_text = all_text + run.text

        print(all_text)

        return all_text

    def Pop_up_message(self,msg_text):
        msg = QMessageBox()
        msg.setWindowTitle("Alert!")
        msg.setText(msg_text)
        x = msg.exec()  # this will show our messagebox
        
    regionList = []
    def Send_button(self):
        

        if self.exl_file_path==None:
            self.Pop_up_message("Please Select Excel File")
            return
        
        if self.output_folder_path==None:
            self.Pop_up_message("Please Select Output Folder")
            return


        if self.body_tem_file_path==None:
            self.Pop_up_message("Please Select Body Template")
            return

        if self.attachment_tem_file_path==None:
            self.Pop_up_message("Please Select Attachment Template")
            return
        with open("subject.html", "w") as file:
            file.write(self.lineEdit.text())
        file.close()
        self.subject_tem_file_path="subject.html"
        
        if self.subject_tem_file_path==None:
            self.Pop_up_message("Please Select subject Template")
            return

        print("Starting to send Emails...")  
        if self.exl_file_path !=None: 
            print(self.exl_file_path)
            data = pd.read_excel(self.exl_file_path)
            col_value=[]
            for name in data:
                col_value.append(name)
                print(name)
            email=""
            for i in col_value:
                if i=='Email' or i=='email' or i=='EMAIL' or str(i).upper() == "EMAIL":
                    email=i
                    # print(email+"##")
        else:
            print("Please Select Excel file first.")
            return
        
        if self.output_folder_path !=None: 
            print("output folder is set.")
            print(self.output_folder_path)

        if self.subject_tem_file_path !=None: 
            print("subject template is set.")
            print(self.subject_tem_file_path)
        else:
            print("Please Select output fodler.")
            return

        user_name = ""
        # getting the names and the emails
        l1 = []

        for index, row in data.iterrows():
            l1.append(row.to_list())
            

        # names = data['Name']
        emails = data[email]
        # salarys=data['Salry']
        # positions=data['Position']

        value = self.comboBox.currentText()


        #here
        for i in range(len(emails)):
            try:
                # for every record get the name and the email addresses
                l2=[]
                for j in l1[i]:
                    j=str(j)
                    if 'nan' in j or 'NaN' in j:
                        j = "unknown"
                    l2.append(j)    
                    
                print(l2)
                # n = names[i]
                # n = str(n)
                # if 'nan' in n or 'NaN' in n:
                #     n = "unknown"

                e = emails[i]
                email=str(e)
                if 'nan' in email or 'NaN' in email:
                    email = "unknown"
                print(email)

                # fn=file_name[i]
                fn = data[value]
                fn = fn[i]

                fn = str(fn)
                if 'nan' in fn or 'NaN' in fn:
                    fn = "unknown"

                #set defalt email
                if 'unknown' in email:
                    email=lines[5].strip()
                    email = email.split("=")[1]
                    
                self.pdf_name_in_mail = fn


                send_to=email
                # the message to be emailed
                row_value=l2
                path1=Document(self.attachment_tem_file_path)
                z,mail_msg,subject_msg = self.ch_var(path1,col_value,row_value)

                # if 'REGION' in [col.upper() for col in self.allColumns]:
                #     print("Region Present")
                region = ''
                for col in self.allColumns:
                    if col.upper() == 'CITY':
                        region = data[col][i]
                        # print("Region Present",region)
                        self.regionList.append(region)
                        
                # name = "output\\"+name
                if region == '':
                    z.save(self.output_folder_path+"//"+fn+"_"+str(i+2)+'.docx')
                    filename = self.output_folder_path+"//"+fn+"_"+str(i+2)+".pdf"
                    convert(self.output_folder_path+"//"+fn+"_"+str(i+2)+".docx", filename)
                    print("pdf generation done")
                else:
                    if not os.path.exists(self.output_folder_path+"/"+region):
                        os.makedirs(self.output_folder_path+"/"+region)
                    z.save(self.output_folder_path+"/"+region+"//"+fn+"_"+str(i+2)+'.docx')
                    filename = self.output_folder_path+"/"+region+"//"+fn+"_"+str(i+2)+".pdf"
                    convert(self.output_folder_path+"/"+region+"//"+fn+"_"+str(i+2)+".docx", filename)
                    print("pdf generation done")
                    
                # z.save(self.output_folder_path+"//"+fn+"_"+str(i+2)+'.docx')
                # filename = self.output_folder_path+"//"+fn+"_"+str(i+2)+".pdf"
                # convert(self.output_folder_path+"//"+fn+"_"+str(i+2)+".docx", filename)
                # print("pdf generation done")

                # self.sendMail_new(filename,send_to,self.body_tem_file_path,mail_msg,subject_msg)

                print("Mail sent: "+email)        
                print("waiting "+self.wait_time +" seconds")
                time.sleep(int(self.wait_time))
                email = ''
                print("************************ Cycle Completed.************************")

            except Exception as e:
                print(e)

        self.regionList = list(dict.fromkeys(self.regionList))
        for region in self.regionList:
            folder = self.output_folder_path+"/"+region
            merge_pdfs_in_directory(folder,region)
        #all variables reset
        self.exl_file_path = None
        self.output_folder_path = None
        self.body_tem_file_path = None
        self.attachment_tem_file_path = None
        self.subject_tem_file_path=None
        self.regionList.clear()
        self.comboBox.clear()

        print("Process completed.")

def merge_pdfs_in_directory(directory, region):
    merger = PdfMerger()

    # Get a list of all PDF files in the directory
    pdf_files = [f for f in os.listdir(directory) if f.endswith('.pdf')]

    # Sort the files alphabetically (optional)
    # pdf_files.sort()

    # Merge the PDF files
    for pdf_file in pdf_files:
        file_path = os.path.join(directory, pdf_file)
        merger.append(file_path)

    # Save the merged PDF file
    # print(region)
    outputPdfName = f'{region}.pdf'
    # print(outputPdfName)
    merged_file_path = os.path.join(directory, outputPdfName)
    merger.write(merged_file_path)
    merger.close()

    print(f'PDF files for region {region} merged successfully.')
    
    
if __name__ == "__main__":
    import sys
    filename = "pandoc-3.1.2-windows-x86_64.msi"
    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)

    if os.path.exists(filepath):
        print(f"The file {filename} exists in the base location of the code.")
    else:
        # print(f"The file {filename} does not exist in the base location of the code.")
        download_pandoc(version='3.1.2')
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    with open('Config_file.txt') as f:
        lines = f.readlines()
    
    title_name=lines[0].strip()
    title_name = title_name.split("=")[1]

    bg_color_set=lines[1].strip()
    bg_color_set = bg_color_set.split("=")[1]

    btn_color_set =  lines[2].strip()
    btn_color_set = btn_color_set.split("=")[1]

    mail_subject=''
    mail_subject = ''

    mail_body=''
    mail_body=''


    all_text_color=lines[6].strip()
    all_text_color=all_text_color.split("=")[1]

    send_button_text_color=lines[7].strip()
    send_button_text_color=send_button_text_color.split("=")[1]

    main_window_title=lines[11].strip()
    main_window_title=main_window_title.split("=")[1]

    main_icon=lines[12].strip()
    main_icon=main_icon.split("=")[1]

    wait_time_in_sec=lines[8].strip()
    wait_time_in_sec=wait_time_in_sec.split("=")[1]
    myLabel= QLabel()
    myLabel.setAutoFillBackground(True) # This is important!!
    

    ui.setupUi(MainWindow,btn_color_set,bg_color_set,title_name,mail_subject,mail_body,all_text_color,send_button_text_color,wait_time_in_sec,main_window_title,main_icon)
    MainWindow.show()
    sys.exit(app.exec())
